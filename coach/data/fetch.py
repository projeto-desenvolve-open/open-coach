import requests
import json
import os
from docx import Document
from docx.shared import Pt
from uuid import uuid4

# Configurações
base_url = "https://projetodesenvolve.online"
client_id = "mycoach-consumer"
client_secret = "0vpaOnw1lzL0iaGSEcGjkUusclqHscYkWPrH1tA9fU48AvX0C9eeWO1iy2Yi3wwjdUilrAQzT3PyXbC2qW2s8rDINDo6dnpNDv2ZiXlGHOjRMp2AuzG8BAmHtpeX0Orb"
username = "admin"
password = "sWqyPA1ihS03"

# Carregar cursos do arquivo cursos.json
try:
    with open('cursos.json', 'r') as file:
        courses = json.load(file)
    print(f"Total de cursos carregados de cursos.json: {len(courses)}")
except FileNotFoundError:
    print("Erro: Arquivo cursos.json não encontrado no diretório atual.")
    exit(1)
except json.JSONDecodeError as e:
    print(f"Erro ao decodificar cursos.json: {e}")
    exit(1)

# Obter token
token_url = f"{base_url}/oauth2/access_token/"
data = {
    "grant_type": "password",
    "client_id": client_id,
    "client_secret": client_secret,
    "username": username,
    "password": password,
    "scope": "read write"
}

response = requests.post(token_url, data=data)
if response.status_code != 200:
    print(f"Erro ao obter token: {response.status_code} - {response.text}")
    exit(1)

try:
    token_data = response.json()
    access_token = token_data["access_token"]
    print(f"Access Token: {access_token}")
except (ValueError, KeyError) as e:
    print(f"Erro ao decodificar JSON do token: {e} - Resposta: {response.text}")
    exit(1)

# Buscar notas de todos os alunos para cada curso
all_grades = {}
for course in courses:
    course_id = course["course_id"]
    course_name = course["course_name"]
    print(f"\nBuscando notas para o curso: {course_name} ({course_id})")

    grades_url = f"{base_url}/api/grades/v1/gradebook/{course_id}/"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    params = {
        "page_size": 100,
        "excluded_course_roles": "all"
    }

    all_results = []
    next_url = grades_url

    while next_url:
        try:
            response = requests.get(next_url, headers=headers, params=params if next_url == grades_url else None, timeout=10)
            if response.status_code != 200:
                print(f"Erro ao buscar notas do curso {course_id}: {response.status_code} - {response.text}")
                all_grades[course_id] = {"course_name": course_name, "error": response.text, "grades": []}
                break

            try:
                data = response.json()
                print(f"Resposta bruta do curso {course_id} (página):", json.dumps(data, indent=2))

                all_results.extend(data["results"])
                next_url = data.get("next")
                print(f"Coletados {len(all_results)} resultados para o curso {course_id} até agora...")
            except (ValueError, KeyError) as e:
                print(f"Erro ao decodificar JSON das notas do curso {course_id}: {e} - Resposta: {response.text}")
                all_grades[course_id] = {"course_name": course_name, "error": f"Erro ao processar resposta: {str(e)}", "grades": []}
                break
        except requests.RequestException as e:
            print(f"Erro de rede ao buscar notas do curso {course_id}: {str(e)}")
            all_grades[course_id] = {"course_name": course_name, "error": f"Erro de rede: {str(e)}", "grades": []}
            break

    # Processar notas para calcular porcentagens
    processed_grades = []
    for result in all_results:
        try:
            total_score = sum(
                section['score_earned'] for section in result['section_breakdown']
                if section['score_possible'] > 0
            )
            total_possible = sum(
                section['score_possible'] for section in result['section_breakdown']
                if section['score_possible'] > 0
            )
            grade = (total_score / total_possible * 100) if total_possible > 0 else 0
            processed_grades.append({
                "user_id": result.get("user_id"),
                "username": result.get("username"),
                "email": result.get("email"),
                "calculated_grade": round(grade, 2),
                "section_breakdown": result.get("section_breakdown")
            })
        except (KeyError, ZeroDivisionError) as e:
            print(f"Erro ao processar notas para aluno {result.get('email', 'desconhecido')} no curso {course_id}: {str(e)}")
            processed_grades.append({
                "user_id": result.get("user_id"),
                "username": result.get("username"),
                "email": result.get("email"),
                "calculated_grade": None,
                "section_breakdown": result.get("section_breakdown"),
                "error": str(e)
            })

    all_grades[course_id] = {
        "course_name": course_name,
        "grades": processed_grades,
        "error": None if processed_grades else "Nenhuma nota encontrada"
    }
    print(f"Total de alunos encontrados no curso {course_id}: {len(processed_grades)}")

# Exibir resumo das notas
print("\nResumo das notas por curso:")
for course_id, grade_info in all_grades.items():
    print(f"\nCurso: {grade_info['course_name']} ({course_id})")
    if grade_info.get("error"):
        print(f"Erro: {grade_info['error']}")
    else:
        print(f"Total de alunos: {len(grade_info['grades'])}")
        for grade in grade_info["grades"]:
            print(f"\nAluno: {grade['email']} ({grade['username']})")
            if grade.get("error"):
                print(f"Erro: {grade['error']}")
            else:
                print(f"Nota Calculada: {grade['calculated_grade']}%")
                print("Detalhes:", json.dumps(grade["section_breakdown"], indent=2))

# Contagem total de alunos
total_alunos = sum(len(grade_info["grades"]) for grade_info in all_grades.values())
print(f"\nTotal de alunos em todos os cursos: {total_alunos}")
successful_courses = sum(1 for g in all_grades.values() if not g.get("error"))
print(f"Total de cursos processados com sucesso: {successful_courses}/{len(courses)}")

# Salvar dados em JSON
try:
    with open('all_grades.json', 'w', encoding='utf-8') as json_file:
        json.dump(all_grades, json_file, ensure_ascii=False, indent=2)
    print("\nArquivo JSON 'all_grades.json' gerado com sucesso.")
except Exception as e:
    print(f"Erro ao salvar JSON: {str(e)}")

# Criar documento DOCX
try:
    doc = Document()
    doc.add_heading('Relatório de Notas por Curso', 0)

    for course_id, grade_info in all_grades.items():
        doc.add_heading(f"Curso: {grade_info['course_name']} ({course_id})", level=1)
        if grade_info.get("error"):
            doc.add_paragraph(f"Erro: {grade_info['error']}")
        else:
            doc.add_paragraph(f"Total de alunos: {len(grade_info['grades'])}")
            for grade in grade_info["grades"]:
                doc.add_heading(f"Aluno: {grade['username']}", level=2)
                if grade['email']:
                    doc.add_paragraph(f"Email: {grade['email']}")
                if grade.get("error"):
                    doc.add_paragraph(f"Erro: {grade['error']}")
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"Nota Calculada: {grade['calculated_grade']}%").bold = True
                    doc.add_paragraph("Detalhes das Atividades:", style='Intense Quote')
                    for section in grade['section_breakdown']:
                        score_text = f"{section['score_earned']}/{section['score_possible']}" if section['score_possible'] > 0 else "Não avaliado"
                        doc.add_paragraph(
                            f"{section['label']} ({section['subsection_name']}): {score_text} "
                            f"({'Tentado' if section['attempted'] else 'Não tentado'})"
                        )

    doc.save('grades_report.docx')
    print("Arquivo DOCX 'grades_report.docx' gerado com sucesso.")
except Exception as e:
    print(f"Erro ao criar DOCX: {str(e)}")